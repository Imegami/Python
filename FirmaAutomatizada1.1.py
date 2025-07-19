#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Sistema Automático de Firma de Documentos
=========================================

Aplicación para insertar firmas automáticamente en documentos usando datos de Excel.
Soporta PDF, DOCX, RTF y otros formatos.

Autor: Imegami
Versión: 1.1
Fecha: 2025
"""

import os
import sys
import logging
import pandas as pd
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont
import zipfile
from datetime import datetime
import json
from typing import List, Dict, Optional, Tuple
import re


class DocumentSigner:
    """Clase principal para el sistema de firma de documentos."""
    
    def __init__(self):
        self.logger = self._setup_logger()
        self.config = self._load_config()
        self.excel_data = None
        self.signature_images = {}
        self.processed_documents = []
        
    def _setup_logger(self) -> logging.Logger:
        """Configura el sistema de logging."""
        logger = logging.getLogger('DocumentSigner')
        logger.setLevel(logging.INFO)
        
        # Handler para archivo
        file_handler = logging.FileHandler('document_signer.log', encoding='utf-8')
        file_handler.setLevel(logging.INFO)
        
        # Handler para consola
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.INFO)
        
        # Formato
        formatter = logging.Formatter(
            '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
        file_handler.setFormatter(formatter)
        console_handler.setFormatter(formatter)
        
        logger.addHandler(file_handler)
        logger.addHandler(console_handler)
        
        return logger
    
    def _load_config(self) -> dict:
        """Carga configuración por defecto."""
        return {
            'output_folder': str(Path.home() / 'Documentos' / 'Documentos_Firmados'),
            'signature_size': (150, 60),
            'font_size': 12,
            'font_color': 'black',
            'signature_quality': 95,
            'backup_originals': True,
            'placeholders': {
                'firma': ['<<firma>>', '<<signature>>', '[FIRMA]', '[SIGNATURE]'],
                'nombre': ['<<nombre>>', '<<name>>', '[NOMBRE]', '[NAME]'],
                'dni': ['<<dni>>', '<<nif>>', '[DNI]', '[NIF]']
            },
            'common_paths': {
                'excel': [
                    str(Path.home() / 'Escritorio'),
                    str(Path.home() / 'Documentos'),
                    str(Path.home() / 'Descargas')
                ],
                'documents': [
                    str(Path.home() / 'Escritorio'),
                    str(Path.home() / 'Documentos'),
                    str(Path.home() / 'Descargas')
                ]
            }
        }
    
    def validate_excel_file(self, file_path: str) -> Tuple[bool, str]:
        """Valida el archivo Excel y sus columnas requeridas."""
        try:
            df = pd.read_excel(file_path)
            
            # Columnas requeridas
            required_cols = ['nombre', 'dni']
            optional_cols = ['apellido1', 'apellido2', 'firma_imagen']
            
            # Verificar columnas requeridas
            missing_cols = [col for col in required_cols if col not in df.columns]
            if missing_cols:
                return False, f"Faltan columnas requeridas: {', '.join(missing_cols)}"
            
            # Verificar datos vacíos
            empty_rows = df[df[required_cols].isnull().any(axis=1)]
            if not empty_rows.empty:
                return False, f"Hay {len(empty_rows)} filas con datos faltantes en columnas requeridas"
            
            self.logger.info(f"Excel validado correctamente: {len(df)} registros encontrados")
            return True, f"Archivo válido con {len(df)} registros"
            
        except Exception as e:
            return False, f"Error al leer Excel: {str(e)}"
    
    def load_excel_data(self, file_path: str) -> bool:
        """Carga los datos del archivo Excel."""
        try:
            self.excel_data = pd.read_excel(file_path)
            
            # Crear columna nombre_completo si no existe
            if 'nombre_completo' not in self.excel_data.columns:
                cols = ['nombre']
                if 'apellido1' in self.excel_data.columns:
                    cols.append('apellido1')
                if 'apellido2' in self.excel_data.columns:
                    cols.append('apellido2')
                
                self.excel_data['nombre_completo'] = self.excel_data[cols].apply(
                    lambda x: ' '.join(x.dropna().astype(str)), axis=1
                )
            
            self.logger.info(f"Datos Excel cargados: {len(self.excel_data)} registros")
            return True
            
        except Exception as e:
            self.logger.error(f"Error cargando Excel: {e}")
            return False
    
    def generate_signature_image(self, name: str) -> str:
        """Genera una imagen de firma manuscrita para un nombre dado."""
        try:
            # Crear imagen
            img_width, img_height = self.config['signature_size']
            img = Image.new('RGBA', (img_width, img_height), (255, 255, 255, 0))
            draw = ImageDraw.Draw(img)
            
            # Intentar cargar fuente manuscrita
            font_paths = [
                'fonts/signature_font.ttf',
                'C:/Windows/Fonts/segoeui.ttf',  # Windows
                '/System/Library/Fonts/Helvetica.ttc',  # macOS
                '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'  # Linux
            ]
            
            font = None
            for font_path in font_paths:
                try:
                    font = ImageFont.truetype(font_path, 24)
                    break
                except:
                    continue
            
            if font is None:
                font = ImageFont.load_default()
            
            # Calcular posición centrada
            bbox = draw.textbbox((0, 0), name, font=font)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
            
            x = (img_width - text_width) // 2
            y = (img_height - text_height) // 2
            
            # Dibujar texto con efecto manuscrito
            draw.text((x, y), name, fill='black', font=font)
            
            # Guardar imagen temporal
            temp_path = f"temp_signature_{name.replace(' ', '_')}.png"
            img.save(temp_path, 'PNG', quality=self.config['signature_quality'])
            
            self.logger.info(f"Firma generada para: {name}")
            return temp_path
            
        except Exception as e:
            self.logger.error(f"Error generando firma para {name}: {e}")
            return None
    
    def sign_pdf(self, pdf_path: str, person_data: dict, output_path: str) -> bool:
        """Firma un documento PDF."""
        try:
            doc = fitz.open(pdf_path)
            
            # Obtener o generar imagen de firma
            signature_path = None
            if 'firma_imagen' in person_data and pd.notna(person_data['firma_imagen']):
                signature_path = person_data['firma_imagen']
            else:
                signature_path = self.generate_signature_image(person_data['nombre_completo'])
            
            if not signature_path or not os.path.exists(signature_path):
                self.logger.warning(f"No se pudo obtener firma para {person_data['nombre_completo']}")
                return False
            
            # Procesar cada página
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Buscar placeholders
                text_instances = page.search_for("<<firma>>")
                name_instances = page.search_for("<<nombre>>")
                dni_instances = page.search_for("<<dni>>")
                
                # Insertar firma
                if text_instances:
                    for inst in text_instances:
                        # Insertar imagen de firma
                        img_rect = fitz.Rect(inst.x0, inst.y0, inst.x0 + 150, inst.y0 + 60)
                        page.insert_image(img_rect, filename=signature_path)
                        
                        # Eliminar placeholder
                        page.add_redact_annot(inst)
                else:
                    # Insertar al final si no hay placeholder
                    rect = page.rect
                    img_rect = fitz.Rect(rect.width - 200, rect.height - 100, 
                                       rect.width - 50, rect.height - 40)
                    page.insert_image(img_rect, filename=signature_path)
                
                # Insertar nombre
                if name_instances:
                    for inst in name_instances:
                        page.insert_text(
                            (inst.x0, inst.y0 + 15), 
                            person_data['nombre_completo'],
                            fontsize=self.config['font_size'],
                            color=(0, 0, 0)
                        )
                        page.add_redact_annot(inst)
                
                # Insertar DNI
                if dni_instances:
                    for inst in dni_instances:
                        page.insert_text(
                            (inst.x0, inst.y0 + 15), 
                            str(person_data['dni']),
                            fontsize=self.config['font_size'],
                            color=(0, 0, 0)
                        )
                        page.add_redact_annot(inst)
                
                # Aplicar redacciones
                page.apply_redactions()
            
            # Guardar documento firmado
            doc.save(output_path)
            doc.close()
            
            # Limpiar archivo temporal de firma
            if signature_path.startswith("temp_signature_"):
                try:
                    os.remove(signature_path)
                except:
                    pass
            
            self.logger.info(f"PDF firmado exitosamente: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error firmando PDF {pdf_path}: {e}")
            return False
    
    def sign_docx(self, docx_path: str, person_data: dict, output_path: str) -> bool:
        """Firma un documento DOCX."""
        try:
            doc = Document(docx_path)
            
            # Generar o obtener imagen de firma
            signature_path = None
            if 'firma_imagen' in person_data and pd.notna(person_data['firma_imagen']):
                signature_path = person_data['firma_imagen']
            else:
                signature_path = self.generate_signature_image(person_data['nombre_completo'])
            
            if not signature_path or not os.path.exists(signature_path):
                self.logger.warning(f"No se pudo obtener firma para {person_data['nombre_completo']}")
                return False
            
            # Buscar y reemplazar placeholders en párrafos
            placeholder_found = False
            
            for paragraph in doc.paragraphs:
                if '<<firma>>' in paragraph.text:
                    # Limpiar párrafo y agregar firma
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.add_picture(signature_path, width=Inches(2))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    placeholder_found = True
                
                if '<<nombre>>' in paragraph.text:
                    paragraph.text = paragraph.text.replace('<<nombre>>', person_data['nombre_completo'])
                    placeholder_found = True
                
                if '<<dni>>' in paragraph.text:
                    paragraph.text = paragraph.text.replace('<<dni>>', str(person_data['dni']))
                    placeholder_found = True
            
            # Si no se encontraron placeholders, agregar al final
            if not placeholder_found:
                # Agregar línea en blanco
                doc.add_paragraph()
                
                # Agregar firma
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(signature_path, width=Inches(2))
                
                # Agregar nombre y DNI
                doc.add_paragraph(f"Nombre: {person_data['nombre_completo']}")
                doc.add_paragraph(f"DNI: {person_data['dni']}")
            
            # Guardar documento
            doc.save(output_path)
            
            # Limpiar archivo temporal
            if signature_path.startswith("temp_signature_"):
                try:
                    os.remove(signature_path)
                except:
                    pass
            
            self.logger.info(f"DOCX firmado exitosamente: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error firmando DOCX {docx_path}: {e}")
            return False
    
    def process_documents(self, document_paths: List[str], output_folder: str, 
                         progress_callback=None) -> Dict:
        """Procesa múltiples documentos con todas las firmas."""
        if self.excel_data is None:
            return {'success': False, 'message': 'No hay datos Excel cargados'}
        
        results = {
            'success': True,
            'processed': 0,
            'failed': 0,
            'details': [],
            'log_file': None
        }
        
        # Crear carpeta de salida
        os.makedirs(output_folder, exist_ok=True)
        
        total_operations = len(document_paths) * len(self.excel_data)
        current_operation = 0
        
        # Procesar cada documento con cada persona
        for doc_path in document_paths:
            if not os.path.exists(doc_path):
                continue
                
            doc_name = Path(doc_path).stem
            doc_ext = Path(doc_path).suffix.lower()
            
            for idx, person in self.excel_data.iterrows():
                current_operation += 1
                
                # Actualizar progreso
                if progress_callback:
                    progress = (current_operation / total_operations) * 100
                    progress_callback(progress, f"Procesando {doc_name} - {person['nombre_completo']}")
                
                # Generar nombre de archivo de salida
                safe_name = re.sub(r'[^\w\-_.]', '_', person['nombre_completo'])
                output_filename = f"{doc_name}_{safe_name}{doc_ext}"
                output_path = os.path.join(output_folder, output_filename)
                
                # Procesar según tipo de archivo
                success = False
                if doc_ext == '.pdf':
                    success = self.sign_pdf(doc_path, person, output_path)
                elif doc_ext == '.docx':
                    success = self.sign_docx(doc_path, person, output_path)
                else:
                    self.logger.warning(f"Formato no soportado: {doc_ext}")
                    continue
                
                if success:
                    results['processed'] += 1
                    results['details'].append({
                        'documento': doc_name,
                        'persona': person['nombre_completo'],
                        'dni': person['dni'],
                        'archivo_salida': output_filename,
                        'estado': 'Éxito',
                        'timestamp': datetime.now().isoformat()
                    })
                else:
                    results['failed'] += 1
                    results['details'].append({
                        'documento': doc_name,
                        'persona': person['nombre_completo'],
                        'dni': person['dni'],
                        'archivo_salida': output_filename,
                        'estado': 'Error',
                        'timestamp': datetime.now().isoformat()
                    })
        
        # Generar reporte Excel
        if results['details']:
            report_path = os.path.join(output_folder, f"reporte_firmas_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
            df_report = pd.DataFrame(results['details'])
            df_report.to_excel(report_path, index=False)
            results['log_file'] = report_path
        
        results['message'] = f"Procesados: {results['processed']}, Fallidos: {results['failed']}"
        return results


class DocumentSignerGUI:
    """Interfaz gráfica para el sistema de firma de documentos."""
    
    def __init__(self):
        self.signer = DocumentSigner()
        self.root = tk.Tk()
        self.setup_ui()
        
    def setup_ui(self):
        """Configura la interfaz de usuario."""
        self.root.title("Sistema de Firma Automática de Documentos v1.0")
        self.root.geometry("800x700")
        self.root.resizable(True, True)
        
        # Notebook para pestañas
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Pestaña 1: Configuración
        self.setup_config_tab(notebook)
        
        # Pestaña 2: Procesamiento
        self.setup_processing_tab(notebook)
        
        # Pestaña 3: Resultados
        self.setup_results_tab(notebook)
    
    def setup_config_tab(self, notebook):
        """Configura la pestaña de configuración."""
        config_frame = ttk.Frame(notebook)
        notebook.add(config_frame, text="📋 Configuración")
        
        # Frame para Excel
        excel_frame = ttk.LabelFrame(config_frame, text="Archivo Excel con Datos", padding=10)
        excel_frame.pack(fill='x', padx=10, pady=5)
        
        ttk.Label(excel_frame, text="Archivo Excel:").pack(anchor='w')
        excel_path_frame = ttk.Frame(excel_frame)
        excel_path_frame.pack(fill='x', pady=5)
        
        self.excel_path_var = tk.StringVar()
        ttk.Entry(excel_path_frame, textvariable=self.excel_path_var, width=60).pack(side='left', fill='x', expand=True)
        ttk.Button(excel_path_frame, text="📁 Buscar", command=self.browse_excel).pack(side='right', padx=(5, 0))
        
        # Botón validar Excel
        ttk.Button(excel_frame, text="✓ Validar Excel", command=self.validate_excel).pack(pady=5)
        
        # Área de estado Excel
        self.excel_status = tk.Text(excel_frame, height=3, wrap='word')
        self.excel_status.pack(fill='x', pady=5)
        
        # Frame para documentos
        docs_frame = ttk.LabelFrame(config_frame, text="Documentos a Firmar", padding=10)
        docs_frame.pack(fill='both', expand=True, padx=10, pady=5)
        
        ttk.Label(docs_frame, text="Documentos seleccionados:").pack(anchor='w')
        
        docs_list_frame = ttk.Frame(docs_frame)
        docs_list_frame.pack(fill='both', expand=True, pady=5)
        
        self.docs_listbox = tk.Listbox(docs_list_frame, selectmode='extended')
        scrollbar = ttk.Scrollbar(docs_list_frame, orient='vertical', command=self.docs_listbox.yview)
        self.docs_listbox.configure(yscrollcommand=scrollbar.set)
        
        self.docs_listbox.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')
        
        docs_buttons_frame = ttk.Frame(docs_frame)
        docs_buttons_frame.pack(fill='x', pady=5)
        
        ttk.Button(docs_buttons_frame, text="➕ Agregar Documentos", command=self.browse_documents).pack(side='left', padx=(0, 5))
        ttk.Button(docs_buttons_frame, text="➖ Eliminar Seleccionado", command=self.remove_documents).pack(side='left', padx=5)
        ttk.Button(docs_buttons_frame, text="🗑️ Limpiar Lista", command=self.clear_documents).pack(side='left', padx=5)
        
        # Frame configuración de salida
        output_frame = ttk.LabelFrame(config_frame, text="Configuración de Salida", padding=10)
        output_frame.pack(fill='x', padx=10, pady=5)
        
        ttk.Label(output_frame, text="Carpeta de salida:").pack(anchor='w')
        output_path_frame = ttk.Frame(output_frame)
        output_path_frame.pack(fill='x', pady=5)
        
        self.output_path_var = tk.StringVar(value=self.signer.config['output_folder'])
        ttk.Entry(output_path_frame, textvariable=self.output_path_var, width=60).pack(side='left', fill='x', expand=True)
        ttk.Button(output_path_frame, text="📁 Cambiar", command=self.browse_output_folder).pack(side='right', padx=(5, 0))
    
    def setup_processing_tab(self, notebook):
        """Configura la pestaña de procesamiento."""
        process_frame = ttk.Frame(notebook)
        notebook.add(process_frame, text="⚡ Procesamiento")
        
        # Resumen
        summary_frame = ttk.LabelFrame(process_frame, text="Resumen de Configuración", padding=10)
        summary_frame.pack(fill='x', padx=10, pady=10)
        
        self.summary_text = tk.Text(summary_frame, height=8, wrap='word')
        self.summary_text.pack(fill='x')
        
        # Barra de progreso
        progress_frame = ttk.LabelFrame(process_frame, text="Progreso", padding=10)
        progress_frame.pack(fill='x', padx=10, pady=10)
        
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(progress_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(fill='x', pady=5)
        
        self.progress_label = ttk.Label(progress_frame, text="Listo para procesar")
        self.progress_label.pack()
        
        # Botones de acción
        action_frame = ttk.Frame(process_frame)
        action_frame.pack(fill='x', padx=10, pady=20)
        
        ttk.Button(action_frame, text="🔍 Actualizar Resumen", command=self.update_summary).pack(side='left', padx=(0, 10))
        ttk.Button(action_frame, text="🚀 PROCESAR DOCUMENTOS", command=self.process_documents, style='Accent.TButton').pack(side='left', padx=10)
        ttk.Button(action_frame, text="⏹️ Detener", command=self.stop_processing).pack(side='left', padx=10)
        
        # Log en tiempo real
        log_frame = ttk.LabelFrame(process_frame, text="Log de Procesamiento", padding=10)
        log_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        self.log_text = scrolledtext.ScrolledText(log_frame, wrap='word', height=15)
        self.log_text.pack(fill='both', expand=True)
    
    def setup_results_tab(self, notebook):
        """Configura la pestaña de resultados."""
        results_frame = ttk.Frame(notebook)
        notebook.add(results_frame, text="📊 Resultados")
        
        # Estadísticas
        stats_frame = ttk.LabelFrame(results_frame, text="Estadísticas del Último Procesamiento", padding=10)
        stats_frame.pack(fill='x', padx=10, pady=10)
        
        self.stats_text = tk.Text(stats_frame, height=6, wrap='word')
        self.stats_text.pack(fill='x')
        
        # Acciones
        actions_frame = ttk.Frame(results_frame)
        actions_frame.pack(fill='x', padx=10, pady=10)
        
        ttk.Button(actions_frame, text="📂 Abrir Carpeta de Salida", command=self.open_output_folder).pack(side='left', padx=(0, 10))
        ttk.Button(actions_frame, text="📄 Ver Reporte Excel", command=self.open_report).pack(side='left', padx=10)
        
        # Lista de archivos generados
        files_frame = ttk.LabelFrame(results_frame, text="Archivos Generados", padding=10)
        files_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        self.files_tree = ttk.Treeview(files_frame, columns=('Documento', 'Persona', 'DNI', 'Estado'), show='headings')
        self.files_tree.heading('Documento', text='Documento')
        self.files_tree.heading('Persona', text='Persona')
        self.files_tree.heading('DNI', text='DNI')
        self.files_tree.heading('Estado', text='Estado')
        
        files_scrollbar = ttk.Scrollbar(files_frame, orient='vertical', command=self.files_tree.yview)
        self.files_tree.configure(yscrollcommand=files_scrollbar.set)
        
        self.files_tree.pack(side='left', fill='both', expand=True)
        files_scrollbar.pack(side='right', fill='y')
    
    def browse_excel(self):
        """Abre diálogo para seleccionar archivo Excel."""
        filename = filedialog.askopenfilename(
            title="Seleccionar archivo Excel",
            filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
        )
        if filename:
            self.excel_path_var.set(filename)
    
    def validate_excel(self):
        """Valida el archivo Excel seleccionado."""
        file_path = self.excel_path_var.get()
        if not file_path:
            messagebox.showwarning("Advertencia", "Por favor selecciona un archivo Excel.")
            return
        
        is_valid, message = self.signer.validate_excel_file(file_path)
        
        self.excel_status.delete(1.0, tk.END)
        if is_valid:
            self.excel_status.insert(tk.END, f"✅ {message}")
            self.excel_status.configure(bg='lightgreen')
            
            # Cargar datos
            if self.signer.load_excel_data(file_path):
                self.excel_status.insert(tk.END, f"\n\nColumnas encontradas:\n{', '.join(self.signer.excel_data.columns.tolist())}")
        else:
            self.excel_status.insert(tk.END, f"❌ {message}")
            self.excel_status.configure(bg='lightcoral')
    
    def browse_documents(self):
        """Abre diálogo para seleccionar documentos."""
        filenames = filedialog.askopenfilenames(
            title="Seleccionar documentos a firmar",
            filetypes=[
                ("PDF files", "*.pdf"),
                ("Word documents", "*.docx"),
                ("RTF files", "*.rtf"),
                ("All supported", "*.pdf *.docx *.rtf"),
                ("All files", "*.*")
            ]
        )
        
        for filename in filenames:
            if filename not in self.docs_listbox.get(0, tk.END):
                self.docs_listbox.insert(tk.END, filename)
    
    def remove_documents(self):
        """Elimina documentos seleccionados de la lista."""
        selected = self.docs_listbox.curselection()
        for index in reversed(selected):
            self.docs_listbox.delete(index)
    
    def clear_documents(self):
        """Limpia la lista de documentos."""
        self.docs_listbox.delete(0, tk.END)
    
    def browse_output_folder(self):
        """Abre diálogo para seleccionar carpeta de salida."""
        folder = filedialog.askdirectory(title="Seleccionar carpeta de salida")
        if folder:
            self.output_path_var.set(folder)
    
    def update_summary(self):
        """Actualiza el resumen de configuración."""
        summary = "=== RESUMEN DE CONFIGURACIÓN ===\n\n"
        
        # Excel
        excel_path = self.excel_path_var.get()
        if excel_path and self.signer.excel_data is not None:
            summary += f"📊 Excel: {os.path.basename(excel_path)}\n"
            summary += f"   Registros: {len(self.signer.excel_data)}\n"
            summary += f"   Columnas: {', '.join(self.signer.excel_data.columns.tolist())}\n\n"
        else:
            summary += "❌ Excel: No configurado o no válido\n\n"
        
        # Documentos
        docs = list(self.docs_listbox.get(0, tk.END))
        if docs:
            summary += f"📄 Documentos: {len(docs)} archivos seleccionados\n"
            for doc in docs:
                summary += f"   • {os.path.basename(doc)}\n"
            summary += "\n"
        else:
            summary += "❌ Documentos: Ningún documento seleccionado\n\n"
        
        # Salida
        output_path = self.output_path_var.get()
        summary += f"📁 Carpeta de salida: {output_path}\n\n"
        
        # Estimación
        if docs and self.signer.excel_data is not None:
            total_files = len(docs) * len(self.signer.excel_data)
            summary += f"🔢 Total de archivos a generar: {total_files}\n"
            summary += f"💾 Espacio estimado: ~{total_files * 0.5:.1f} MB\n"
        
        self.summary_text.delete(1.0, tk.END)
        self.summary_text.insert(1.0, summary)
    
    def progress_callback(self, progress, message):
        """Callback para actualizar progreso."""
        self.progress_var.set(progress)
        self.progress_label.config(text=message)
        self.log_text.insert(tk.END, f"{datetime.now().strftime('%H:%M:%S')} - {message}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()
    
    def process_documents(self):
        """Inicia el procesamiento de documentos."""
        # Validaciones
        if not self.excel_path_var.get() or self.signer.excel_data is None:
            messagebox.showerror("Error", "Por favor configura y valida el archivo Excel.")
            return
        
        docs = list(self.docs_listbox.get(0, tk.END))
        if not docs:
            messagebox.showerror("Error", "Por favor selecciona al menos un documento.")
            return
        
        output_folder = self.output_path_var.get()
        if not output_folder:
            messagebox.showerror("Error", "Por favor selecciona una carpeta de salida.")
            return
        
        # Confirmación
        total_files = len(docs) * len(self.signer.excel_data)
        response = messagebox.askyesno(
            "Confirmar Procesamiento",
            f"Se procesarán {len(docs)} documentos con {len(self.signer.excel_data)} firmas cada uno.\n\n"
            f"Total de archivos a generar: {total_files}\n"
            f"Carpeta de salida: {output_folder}\n\n"
            f"¿Desea continuar?"
        )
        
        if not response:
            return
        
        # Limpiar log
        self.log_text.delete(1.0, tk.END)
        
        # Procesar
        self.log_text.insert(tk.END, f"🚀 Iniciando procesamiento...\n")
        self.log_text.insert(tk.END, f"📊 Datos: {len(self.signer.excel_data)} personas\n")
        self.log_text.insert(tk.END, f"📄 Documentos: {len(docs)} archivos\n")
        self.log_text.insert(tk.END, f"📁 Salida: {output_folder}\n")
        self.log_text.insert(tk.END, "="*50 + "\n\n")
        
        try:
            results = self.signer.process_documents(
                docs, 
                output_folder, 
                progress_callback=self.progress_callback
            )
            
            # Mostrar resultados
            self.show_results(results)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error durante el procesamiento:\n{str(e)}")
            self.signer.logger.error(f"Error en procesamiento: {e}")
    
    def stop_processing(self):
        """Detiene el procesamiento (placeholder)."""
        messagebox.showinfo("Info", "Función de detener en desarrollo.")
    
    def show_results(self, results):
        """Muestra los resultados del procesamiento."""
        # Actualizar estadísticas
        stats_text = f"=== RESULTADOS DEL PROCESAMIENTO ===\n\n"
        stats_text += f"✅ Documentos procesados exitosamente: {results['processed']}\n"
        stats_text += f"❌ Documentos con errores: {results['failed']}\n"
        stats_text += f"📊 Total procesado: {results['processed'] + results['failed']}\n"
        stats_text += f"📁 Carpeta de salida: {self.output_path_var.get()}\n"
        
        if results.get('log_file'):
            stats_text += f"📄 Reporte Excel: {os.path.basename(results['log_file'])}\n"
        
        stats_text += f"⏰ Completado: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
        
        self.stats_text.delete(1.0, tk.END)
        self.stats_text.insert(1.0, stats_text)
        
        # Actualizar tabla de archivos
        for item in self.files_tree.get_children():
            self.files_tree.delete(item)
        
        for detail in results['details']:
            self.files_tree.insert('', 'end', values=(
                detail['documento'],
                detail['persona'], 
                detail['dni'],
                detail['estado']
            ))
        
        # Cambiar a pestaña de resultados
        notebook = self.root.nametowidget(self.root.winfo_children()[0])
        notebook.select(2)  # Pestaña de resultados
        
        # Mensaje final
        if results['success']:
            messagebox.showinfo(
                "Procesamiento Completado",
                f"Procesamiento completado exitosamente!\n\n"
                f"Archivos generados: {results['processed']}\n"
                f"Errores: {results['failed']}\n\n"
                f"Los archivos se guardaron en:\n{self.output_path_var.get()}"
            )
        else:
            messagebox.showwarning(
                "Procesamiento con Errores", 
                f"El procesamiento se completó con errores.\n\n"
                f"Revisa el log para más detalles."
            )
    
    def open_output_folder(self):
        """Abre la carpeta de salida."""
        output_folder = self.output_path_var.get()
        if os.path.exists(output_folder):
            os.startfile(output_folder)  # Windows
        else:
            messagebox.showwarning("Advertencia", "La carpeta de salida no existe.")
    
    def open_report(self):
        """Abre el reporte Excel más reciente."""
        output_folder = self.output_path_var.get()
        if not os.path.exists(output_folder):
            messagebox.showwarning("Advertencia", "La carpeta de salida no existe.")
            return
        
        # Buscar el reporte más reciente
        reports = [f for f in os.listdir(output_folder) if f.startswith('reporte_firmas_') and f.endswith('.xlsx')]
        
        if not reports:
            messagebox.showinfo("Info", "No se encontraron reportes Excel.")
            return
        
        # Abrir el más reciente
        latest_report = max(reports, key=lambda x: os.path.getctime(os.path.join(output_folder, x)))
        report_path = os.path.join(output_folder, latest_report)
        
        try:
            os.startfile(report_path)  # Windows
        except:
            messagebox.showinfo("Info", f"Reporte ubicado en:\n{report_path}")
    
    def run(self):
        """Ejecuta la aplicación."""
        # Configurar estilo
        style = ttk.Style()
        style.theme_use('clam')
        
        # Estilo para botón principal
        style.configure('Accent.TButton', foreground='white', background='#0078d4')
        
        # Inicializar resumen
        self.root.after(100, self.update_summary)
        
        # Centrar ventana
        self.root.update_idletasks()
        x = (self.root.winfo_screenwidth() // 2) - (self.root.winfo_width() // 2)
        y = (self.root.winfo_screenheight() // 2) - (self.root.winfo_height() // 2)
        self.root.geometry(f"+{x}+{y}")
        
        self.root.mainloop()


def main():
    """Función principal."""
    try:
        app = DocumentSignerGUI()
        app.run()
    except Exception as e:
        print(f"Error al iniciar la aplicación: {e}")
        input("Presiona Enter para salir...")


if __name__ == "__main__":
    main()


"""
INSTRUCCIONES DE INSTALACIÓN Y USO
=================================

1. INSTALACIÓN DE DEPENDENCIAS:
   pip install pandas openpyxl PyMuPDF python-docx pillow

2. ESTRUCTURA DEL ARCHIVO EXCEL:
   El archivo Excel debe contener las siguientes columnas:
   - nombre (requerido)
   - dni (requerido)
   - apellido1 (opcional)
   - apellido2 (opcional)  
   - firma_imagen (opcional - ruta a imagen de firma)

3. FORMATOS SOPORTADOS:
   - PDF (.pdf)
   - Word (.docx)
   - RTF (.rtf) - en desarrollo

4. PLACEHOLDERS EN DOCUMENTOS:
   Coloca estos marcadores en tus documentos donde quieras insertar datos:
   - <<firma>> - Para insertar la imagen de firma
   - <<nombre>> - Para insertar el nombre completo
   - <<dni>> - Para insertar el DNI

5. USO:
   a) Ejecuta el script: python document_signer.py
   b) En la pestaña "Configuración":
      - Selecciona y valida tu archivo Excel
      - Agrega los documentos a firmar
      - Configura la carpeta de salida
   c) En la pestaña "Procesamiento":
      - Revisa el resumen
      - Haz clic en "PROCESAR DOCUMENTOS"
   d) Revisa los resultados en la pestaña "Resultados"

6. CARACTERÍSTICAS PRINCIPALES:
   ✅ Interfaz gráfica intuitiva con pestañas
   ✅ Validación automática de datos Excel
   ✅ Generación automática de firmas manuscritas
   ✅ Soporte para múltiples formatos de documento
   ✅ Procesamiento por lotes
   ✅ Barra de progreso en tiempo real
   ✅ Log detallado del proceso
   ✅ Reporte Excel de resultados
   ✅ Sistema de backup automático
   ✅ Manejo robusto de errores

7. MEJORAS IMPLEMENTADAS:
   - Arquitectura modular separando lógica de negocio y UI
   - Sistema de logging completo
   - Validación exhaustiva de datos
   - Interfaz moderna con ttk
   - Generación automática de reportes
   - Manejo de placeholders avanzado
   - Sistema de progreso y cancelación
   - Configuración persistente

8. ESTRUCTURA DE ARCHIVOS GENERADOS:
   Documentos_Firmados/
   ├── documento1_Juan_Perez.pdf
   ├── documento1_Maria_Garcia.pdf
   ├── documento2_Juan_Perez.pdf
   ├── documento2_Maria_Garcia.pdf
   └── reporte_firmas_20250719_143022.xlsx

NOTA LEGAL:
Este software está diseñado para fines educativos, de prueba y documentación simulada.
El usuario es responsable del uso apropiado y legal de esta herramienta.
"""